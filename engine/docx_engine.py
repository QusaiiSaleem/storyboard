"""
DOCX Template Engine for Storyboard Generator
==============================================

Builds production-ready DOCX storyboard documents from scratch, matching the
exact visual design of the original template files. Subagents import this
module and call builder methods to create documents without touching binary
template files.

Architecture
------------
    DocxBuilder (base class)
        - Shared page setup, metadata table, RTL helpers
        |
        +-- ObjectivesBuilder   (Group A: infographic-style content table)
        +-- SummaryBuilder      (Group A: same structure as Objectives)
        +-- InfographicBuilder  (Group A: same structure as Objectives)
        +-- DiscussionBuilder   (Group B: card-style content table)
        +-- AssignmentBuilder   (Group B: same structure as Discussion)
        +-- TestBuilder         (Group C: 3 tables - metadata + info + questions)
        +-- ActivityBuilder     (Group C: metadata + N scene tables)
        +-- VideoBuilder        (Group C: 6-row metadata + N 4-col scene tables)

Quick Start
-----------
    from engine.docx_engine import TestBuilder

    builder = TestBuilder(
        project_code="DSAI",
        unit_number=1,
        unit_name="المهارات الرقمية",
        project_name="تطوير 15 مقرر إلكتروني - جامعة نجران",
        institution="جامعة نجران",
        designer="أحمد",
    )
    builder.set_element_name("الاختبار القبلي")
    builder.set_test_info(
        description="الاختبار القبلي للوحدة الأولى",
        instructions="المحاولات المتاحة: محاولة واحدة"
    )
    builder.add_question(
        question_text="ما هو الذكاء الاصطناعي؟",
        choices="أ) ...\nب) ...\nج) ...\nد) ...",
        correct_answer="ج",
        image_description=""
    )
    builder.save("output/DSAI/U01/DSAI_U01_Pre_Test.docx")

Specs Source
------------
All measurements, colors, fonts, and structures come from the template
analysis at: docs/template_analysis.md and docs/template_analysis_raw.json
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Import RTL workaround functions from the shared helpers module.
# These handle the XML-level manipulations that python-docx doesn't
# natively support for Arabic RTL text.
# See: docs/research_findings.md for the full rationale.
from engine.rtl_helpers import (
    docx_set_paragraph_rtl,
    docx_set_run_rtl,
    docx_set_table_rtl,
    docx_set_cell_shading,
    docx_set_cell_borders,
)


# =============================================================================
# CONSTANTS — exact values from template analysis
# =============================================================================

# Page layout (A4 Landscape, all templates)
PAGE_WIDTH_CM = 29.7
PAGE_HEIGHT_CM = 21.0
MARGIN_CM = 2.54  # 1 inch on all sides
HEADER_DISTANCE_CM = 0.25
HEADER_DISTANCE_CM_LARGE = 1.0  # Activity and Video templates use 1.0cm header distance
FOOTER_DISTANCE_CM = 0.7

# Color palette (hex strings WITHOUT the # prefix, for RGBColor)
COLOR_HEADER_BG = "31849B"       # Teal — header row background
COLOR_LABEL_BG = "DBE5F1"        # Light blue-gray — label cell background
COLOR_WHITE = "FFFFFF"            # White — value cells / invisible borders
COLOR_BLACK = "000000"            # Black — text and outer borders
COLOR_HEADER_TEXT = "FFFFFF"      # White — text on teal header
COLOR_GREEN = "007A37"            # Green — page header watermark
COLOR_RED = "FF0000"              # Red — error/highlight (Activity template)
COLOR_VIDEO_SCENE = "CFE2F3"     # Light blue — Video scene headers

# Fonts
FONT_BODY = "Sakkal Majalla"     # Main Arabic content font
FONT_HEADER = "Helvetica Neue"   # Page header font
FONT_FOOTER = "Tahoma"           # Page footer font

# Font sizes (points)
FONT_SIZE_BODY = 12    # Default body text (template docDefaults sz=24 half-pts = 12pt)
FONT_SIZE_HEADER = 14  # Larger headers (tests, video)
FONT_SIZE_FOOTER = 8   # Footer text
FONT_SIZE_PAGE_HEADER = 10  # Page header

# Table widths (in DXA / twips — 1 inch = 1440 dxa)
META_TABLE_WIDTH = 13950       # Standard metadata table
META_COL0_WIDTH = 4050         # Label column
META_COL1_WIDTH = 9900         # Value column

# Group A content table (Objectives, Summary, Infographic)
GROUP_A_TABLE_WIDTH = 14175
GROUP_A_COL0_WIDTH = 3015
GROUP_A_COL1_WIDTH = 11160

# Group B content table (Discussion, Assignment)
GROUP_B_TABLE_WIDTH = 13950
GROUP_B_COL0_WIDTH = 3330
GROUP_B_COL1_WIDTH = 10620

# Test tables
TEST_INFO_TABLE_WIDTH = 14175
TEST_INFO_COL0_WIDTH = 3018
TEST_INFO_COL1_WIDTH = 11157
TEST_Q_TABLE_WIDTH = 14678
TEST_Q_COL_WIDTHS = [3240, 4433, 4050, 2955]

# Activity table
ACTIVITY_TABLE_WIDTH = 13950
ACTIVITY_COL0_WIDTH = 4050
ACTIVITY_COL1_WIDTH = 9900

# Video scene table
VIDEO_TABLE_WIDTH = 13960
VIDEO_COL_WIDTHS = [3490, 3002, 4418, 3050]


# =============================================================================
# HELPER FUNCTIONS — low-level XML manipulation
# =============================================================================

def _set_cell_shading(cell, hex_color):
    """
    Set the background/fill color of a table cell.

    Delegates to docx_set_cell_shading() from rtl_helpers.py which
    creates a NEW XML element each time (critical: XML elements MOVE
    when appended to a different parent, they don't copy).

    Args:
        cell: A python-docx table cell object.
        hex_color: 6-character hex color string (e.g. "31849B"). No # prefix.
    """
    docx_set_cell_shading(cell, hex_color)


def _set_cell_width(cell, width_dxa):
    """
    Set exact cell width in DXA (twips).

    Args:
        cell: A python-docx table cell object.
        width_dxa: Width in DXA units (1 inch = 1440 dxa).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
        tcPr.append(tcW)
    else:
        tcW.set(qn('w:w'), str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')


def _set_cell_vertical_alignment(cell, alignment):
    """
    Set vertical alignment of a cell (top, center, bottom).

    Args:
        cell: A python-docx table cell object.
        alignment: One of "top", "center", "bottom".
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{alignment}"/>')
        tcPr.append(vAlign)
    else:
        vAlign.set(qn('w:val'), alignment)


def _merge_cells_in_row(table, row_idx, start_col, end_col):
    """
    Merge cells horizontally within a single row.

    Args:
        table: python-docx Table object.
        row_idx: The row index (0-based).
        start_col: Starting column index (inclusive).
        end_col: Ending column index (inclusive).
    """
    cell_start = table.cell(row_idx, start_col)
    cell_end = table.cell(row_idx, end_col)
    cell_start.merge(cell_end)


def _set_table_bidi(table):
    """
    Set the table to RTL (BiDi visual) direction.

    This makes the table render right-to-left, which is required for
    Arabic content. Column 0 appears on the RIGHT side of the table.

    Also sets autofit=False so Word respects our explicit column widths
    (research finding: without this, Word may auto-adjust and ignore widths).

    Delegates to docx_set_table_rtl() from rtl_helpers.py which uses
    WD_TABLE_DIRECTION.RTL with an XML fallback.
    """
    docx_set_table_rtl(table)


def _set_table_width(table, width_dxa):
    """
    Set total table width in DXA.

    Args:
        table: python-docx Table object.
        width_dxa: Width in DXA units.
    """
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(width_dxa))
        tblW.set(qn('w:type'), 'dxa')


def _set_table_borders(table, outer_sz=4, inner_sz=18, outer_color="000000",
                        inner_color="FFFFFF"):
    """
    Set table borders.

    The standard template pattern uses:
    - Thin black outer borders (sz=4) for the table frame
    - Thick white inner borders (sz=18) to create invisible separators
      between rows (gives a "card-like" look)

    Some templates override these (e.g., visible inner borders, thicker outer).

    Args:
        table: python-docx Table object.
        outer_sz: Border size for top/bottom/left/right (in half-points).
        inner_sz: Border size for insideH/insideV.
        outer_color: Hex color for outer borders.
        inner_color: Hex color for inner borders.
    """
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)

    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>
        <w:bottom w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>
        <w:left w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>
        <w:right w:val="single" w:sz="{outer_sz}" w:space="0" w:color="{outer_color}"/>
        <w:insideH w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_color}"/>
        <w:insideV w:val="single" w:sz="{inner_sz}" w:space="0" w:color="{inner_color}"/>
    </w:tblBorders>
    """
    tblPr.append(parse_xml(borders_xml))


def _set_metadata_cell_borders(table, num_rows):
    """
    Apply cell-level border overrides to the metadata table.

    Template uses cell-level borders that override table-level defaults:
    - Row 0 (header): top=12, left=12, bottom=8, right=12 (thick frame)
    - Rows 1+  (data): all sides sz=8 (medium frame)

    Args:
        table: python-docx Table object (metadata table).
        num_rows: Number of rows in the table.
    """
    border_thick = {"sz": "12", "val": "single", "color": "000000"}
    border_medium = {"sz": "8", "val": "single", "color": "000000"}

    for row_idx in range(num_rows):
        for col_idx in range(len(table.columns)):
            try:
                cell = table.cell(row_idx, col_idx)
            except Exception:
                continue

            if row_idx == 0:
                # Header row: thick top/left/right, medium bottom
                docx_set_cell_borders(
                    cell,
                    top=border_thick, left=border_thick,
                    right=border_thick, bottom=border_medium,
                )
            else:
                # Data rows: medium on all sides
                docx_set_cell_borders(
                    cell,
                    top=border_medium, bottom=border_medium,
                    left=border_medium, right=border_medium,
                )


def _set_row_height(table, row_idx, height_twips):
    """
    Set explicit row height on a table row.

    Args:
        table: python-docx Table object.
        row_idx: Row index (0-based).
        height_twips: Height in twips.
    """
    row = table.rows[row_idx]
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{height_twips}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def _set_table_indent(table, indent_dxa):
    """
    Set table indent (tblInd) in DXA.

    Used for negative indent on the questions table to extend into margins.

    Args:
        table: python-docx Table object.
        indent_dxa: Indent in DXA (negative values shift left).
    """
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)

    # Remove existing tblInd if any
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)

    tblInd = parse_xml(
        f'<w:tblInd {nsdecls("w")} w:w="{indent_dxa}" w:type="dxa"/>'
    )
    tblPr.append(tblInd)


def _set_paragraph_spacing(paragraph, before=None, after=None, line=None, line_rule=None):
    """
    Set paragraph spacing (before and/or after) in twips, and optionally line spacing.

    Args:
        paragraph: python-docx Paragraph object.
        before: Spacing before in twips (e.g. 240).
        after: Spacing after in twips (e.g. 240).
        line: Line spacing value in twips (e.g. 276 for 1.15x, 360 for 1.5x at 12pt).
              For proportional spacing, use 240 = single, 276 = 1.15, 360 = 1.5, 480 = double.
        line_rule: Line spacing rule — "auto" for proportional (default), "exact" for fixed.
    """
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))
    if line is not None:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), line_rule or "auto")


def _set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """
    Set cell margins (internal padding) in DXA (twips).

    This controls the space between the cell border and the text inside.
    57 dxa ~ 0.1cm, 113 dxa ~ 0.2cm.

    Args:
        cell: python-docx table cell.
        top: Top margin in DXA.
        bottom: Bottom margin in DXA.
        left: Left margin in DXA.
        right: Right margin in DXA.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
        tcPr.append(tcMar)
    for side, value in [('top', top), ('bottom', bottom), ('start', right), ('end', left)]:
        # Note: 'start' maps to right in RTL, 'end' maps to left in RTL
        if value is not None:
            el = tcMar.find(qn(f'w:{side}'))
            if el is None:
                el = parse_xml(f'<w:{side} {nsdecls("w")} w:w="{value}" w:type="dxa"/>')
                tcMar.append(el)
            else:
                el.set(qn('w:w'), str(value))
                el.set(qn('w:type'), 'dxa')


def _set_paragraph_bidi(paragraph):
    """
    Set a paragraph to RTL (BiDi) direction for Arabic text.

    Delegates to docx_set_paragraph_rtl() from rtl_helpers.py which uses
    insert_element_before() to place <w:bidi/> in the correct XML schema
    position within <w:pPr>. This is more robust than a simple append().
    """
    docx_set_paragraph_rtl(paragraph)


def _add_rtl_run(paragraph, text, font_name=FONT_BODY, font_size_pt=None,
                  bold=False, color_hex=None):
    """
    Add a text run with Arabic/RTL formatting to a paragraph.

    This sets both the regular font and the complex-script (cs) font to
    the same value, which is required for Arabic text to render correctly.

    Uses rtl_helpers.docx_set_run_rtl() for the <w:rtl/> element and
    manually sets w:rFonts for cs/ascii/hAnsi fonts plus w:szCs for the
    complex-script font size.

    IMPORTANT (from research_findings.md):
    - Always use add_run() — never set paragraph.text directly
    - font.name alone is ignored for RTL text; must also set w:rFonts w:cs
    - w:szCs must match w:sz for Arabic text sizing to work

    Args:
        paragraph: python-docx Paragraph object.
        text: The Arabic text content.
        font_name: Font family name (default: Sakkal Majalla).
        font_size_pt: Font size in points (None = use default).
        bold: Whether the text is bold.
        color_hex: 6-char hex color (e.g. "000000"). None = inherit.

    Returns:
        The created Run object.
    """
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.bold = bold

    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)

    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)

    # Set complex-script font properties (required for Arabic rendering).
    # font.name alone is IGNORED when RTL is set — must set w:rFonts w:cs
    # separately via XML. See: research_findings.md section 3.6
    rPr = run._r.get_or_add_rPr()

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    # Set complex-script font size (w:szCs) — must match w:sz
    if font_size_pt is not None:
        szCs = rPr.find(qn('w:szCs'))
        half_points = str(int(font_size_pt * 2))
        if szCs is None:
            szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{half_points}"/>')
            rPr.append(szCs)
        else:
            szCs.set(qn('w:val'), half_points)

    # Set <w:rtl/> on the run for proper CS font selection
    docx_set_run_rtl(run)

    return run


def _write_cell(cell, text, font_name=FONT_BODY, font_size_pt=None,
                bold=False, color_hex=None, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                rtl=True, shading_hex=None, vertical_alignment="center",
                line_spacing=276, space_after=120):
    """
    Write text into a cell with full Arabic formatting.

    This is the main workhorse function. It clears the cell, writes text
    with proper RTL formatting, and applies shading/alignment.

    Args:
        cell: python-docx table cell.
        text: Arabic text to write.
        font_name: Font family (default: Sakkal Majalla).
        font_size_pt: Font size in points (None = use FONT_SIZE_BODY=12).
        bold: Bold text.
        color_hex: Text color hex (e.g. "000000").
        alignment: Paragraph alignment (default RIGHT for Arabic).
        rtl: Enable RTL paragraph direction.
        shading_hex: Cell background color hex (e.g. "DBE5F1").
        vertical_alignment: "center", "top", or "bottom".
        line_spacing: Line spacing in twips (276=1.15x, 360=1.5x). None to skip.
        space_after: Space after paragraph in twips (120=~Pt(6)). None to skip.
    """
    # Default font size to FONT_SIZE_BODY (12pt) if not specified
    if font_size_pt is None:
        font_size_pt = FONT_SIZE_BODY

    # Clear existing content
    for p in cell.paragraphs:
        p.clear()

    # Use first existing paragraph
    para = cell.paragraphs[0]
    para.alignment = alignment

    if rtl:
        _set_paragraph_bidi(para)

    _add_rtl_run(para, text, font_name=font_name, font_size_pt=font_size_pt,
                 bold=bold, color_hex=color_hex)

    # Apply line spacing and paragraph spacing for readability
    if line_spacing is not None or space_after is not None:
        _set_paragraph_spacing(para, line=line_spacing, after=space_after)

    if shading_hex:
        _set_cell_shading(cell, shading_hex)

    if vertical_alignment:
        _set_cell_vertical_alignment(cell, vertical_alignment)

    # Apply default cell margins (padding) for text breathing room
    # 85 dxa ~ 0.15cm on each side
    _set_cell_margins(cell, top=57, bottom=57, left=85, right=85)


# =============================================================================
# BASE CLASS — DocxBuilder
# =============================================================================

class DocxBuilder:
    """
    Base builder for all DOCX storyboard templates.

    Handles the shared setup that every template needs:
    - A4 Landscape page with 1-inch margins
    - The standard 7-row metadata table (project info)
    - RTL text direction
    - Page footer with page numbers
    - Save to disk

    ASCII mockup of the metadata table (viewed in RTL — col 0 is on the RIGHT):
    +============================================+
    |  قالب سيناريو ...  (teal bg, white text)   |  <- Row 0 (merged, title)
    +--------------------+-----------------------+
    |     رمز العنصر     | DSAI_U01_Pre_Test     |  <- Row 1
    +--------------------+-----------------------+
    |    اسم المشروع     | تطوير 15 مقرر ...     |  <- Row 2
    +--------------------+-----------------------+
    |  رقم/اسم الوحدة    | الوحدة الأولى          |  <- Row 3
    +--------------------+-----------------------+
    |    اسم العنصر      | الاختبار القبلي        |  <- Row 4
    +--------------------+-----------------------+
    |  المصمم التعليمي    | أحمد                  |  <- Row 5
    +--------------------+-----------------------+
    |      التاريخ        | 2026-02-08            |  <- Row 6
    +--------------------+-----------------------+

    NOTE: In the actual document, this table is RTL so col 0 (labels)
    appears on the RIGHT and col 1 (values) on the LEFT.

    Subclasses override `build_content()` to add template-specific tables.
    """

    # Subclasses MUST set this to the template header title
    TEMPLATE_TITLE = "قالب سيناريو"

    def __init__(self, project_code, unit_number, unit_name, project_name,
                 institution, designer, logo_left_path=None, logo_right_path=None):
        """
        Initialize the builder with project metadata.

        Args:
            project_code: Short code like "DSAI".
            unit_number: Integer unit number (e.g. 1).
            unit_name: Arabic unit name (e.g. "المهارات الرقمية").
            project_name: Full project name in Arabic.
            institution: Institution/client name in Arabic.
            designer: Designer name in Arabic.
            logo_left_path: Path to eduArabia logo image (left side of header).
            logo_right_path: Path to client logo image (right side of header).

        Example:
            builder = TestBuilder(
                project_code="DSAI",
                unit_number=1,
                unit_name="المهارات الرقمية",
                project_name="تطوير 15 مقرر إلكتروني - جامعة نجران",
                institution="جامعة نجران",
                designer="أحمد محمد",
                logo_left_path="projects/DSAI/branding/eduarabia.png",
                logo_right_path="projects/DSAI/branding/client_logo.png",
            )
        """
        self.project_code = project_code
        self.unit_number = unit_number
        self.unit_name = unit_name
        self.project_name = project_name
        self.institution = institution
        self.designer = designer
        self.logo_left_path = logo_left_path
        self.logo_right_path = logo_right_path
        self.element_name = ""
        self.element_code = ""
        self.date_str = date.today().strftime("%Y-%m-%d")

        # Create the document with A4 Landscape setup
        self.doc = Document()
        self._setup_page()

    def set_element_name(self, name):
        """
        Set the element/storyboard name (e.g. "الاختبار القبلي").

        Args:
            name: Arabic element name string.
        """
        self.element_name = name

    def set_element_code(self, code):
        """
        Set the element code (e.g. "DSAI_U01_Pre_Test").

        If not called, the code is auto-generated from project_code and
        unit_number in the metadata table.

        Args:
            code: Element code string.
        """
        self.element_code = code

    def set_date(self, date_str):
        """
        Set the document date (default: today).

        Args:
            date_str: Date string like "2026-02-08".
        """
        self.date_str = date_str

    def _setup_page(self):
        """
        Configure the document page to A4 Landscape with 1-inch margins.

        Specs from template analysis:
        - Page: 29.7cm x 21.0cm (A4 Landscape)
        - Margins: 2.54cm (1 inch) on all sides
        - Header distance: 0.25cm
        - Footer distance: 0.7cm
        """
        section = self.doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        section.header_distance = Cm(HEADER_DISTANCE_CM)
        section.footer_distance = Cm(FOOTER_DISTANCE_CM)

    def _build_header(self):
        """
        Add page header with logos: eduArabia (left) + client logo (right).

        Template structure: two logos side-by-side in a header table.
        Uses a 1-row, 2-column invisible table to position logos.
        The left cell holds the eduArabia logo, the right cell holds
        the client logo.

        Only adds header if at least one logo path is provided.
        """
        if not self.logo_left_path and not self.logo_right_path:
            return

        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # Use a 1x2 table for logo positioning (no borders)
        header_table = header.add_table(rows=1, cols=2, width=Cm(24))
        header_table.autofit = False

        # Remove all borders from the header table
        _set_table_borders(header_table, outer_sz=0, inner_sz=0,
                           outer_color="FFFFFF", inner_color="FFFFFF")

        cell_left = header_table.cell(0, 0)
        cell_right = header_table.cell(0, 1)

        # Set widths: roughly equal halves
        _set_cell_width(cell_left, 7000)
        _set_cell_width(cell_right, 7000)

        # Left cell: eduArabia logo (left-aligned)
        if self.logo_left_path and os.path.exists(self.logo_left_path):
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_left = p_left.add_run()
            # Template size: cx=1990090 cy=402590 EMU (~2.2in x 0.4in)
            run_left.add_picture(self.logo_left_path, width=Emu(1990090), height=Emu(402590))

        # Right cell: client logo (right-aligned)
        if self.logo_right_path and os.path.exists(self.logo_right_path):
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_right = p_right.add_run()
            # Template size: cx=1073150 cy=832976 EMU (~1.2in x 0.9in)
            run_right.add_picture(self.logo_right_path, width=Emu(1073150), height=Emu(832976))

        # Remove the default empty paragraph that headers start with
        if len(header.paragraphs) > 0:
            first_p = header.paragraphs[0]
            if not first_p.text and first_p != cell_left.paragraphs[0]:
                # Clear the empty default paragraph to avoid extra spacing
                first_p.clear()

    def _add_footer(self):
        """
        Add a page footer with "Page X of Y" in Tahoma 8pt.

        Matches the template footer: Tahoma font, 8pt, black text,
        with tab stops and page number fields.
        """
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()

        # Add top border separator line above the footer text
        # Template has: top border single, sz=4, color=#000000, space=1
        pPr = para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="1" w:color="{COLOR_BLACK}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Add spacing before footer text (template uses before=240)
        _set_paragraph_spacing(para, before=240)

        # Add "Page " text
        run_page = para.add_run("Page ")
        run_page.font.name = FONT_FOOTER
        run_page.font.size = Pt(FONT_SIZE_FOOTER)
        run_page.font.color.rgb = RGBColor.from_string(COLOR_BLACK)
        rPr = run_page._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
        rFonts.set(qn('w:cs'), FONT_FOOTER)
        rFonts.set(qn('w:ascii'), FONT_FOOTER)
        rFonts.set(qn('w:hAnsi'), FONT_FOOTER)

        # Add PAGE field
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_fld1 = para.add_run()
        run_fld1._r.append(fldChar_begin)
        run_fld1.font.name = FONT_FOOTER
        run_fld1.font.size = Pt(FONT_SIZE_FOOTER)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_instr = para.add_run()
        run_instr._r.append(instrText)

        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_fld2 = para.add_run()
        run_fld2._r.append(fldChar_end)

        # Add " of " text
        run_of = para.add_run(" of ")
        run_of.font.name = FONT_FOOTER
        run_of.font.size = Pt(FONT_SIZE_FOOTER)
        run_of.font.color.rgb = RGBColor.from_string(COLOR_BLACK)

        # Add NUMPAGES field
        fldChar_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_fld3 = para.add_run()
        run_fld3._r.append(fldChar_begin2)
        run_fld3.font.name = FONT_FOOTER
        run_fld3.font.size = Pt(FONT_SIZE_FOOTER)

        instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run_instr2 = para.add_run()
        run_instr2._r.append(instrText2)

        fldChar_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_fld4 = para.add_run()
        run_fld4._r.append(fldChar_end2)

    def create_metadata_table(self):
        """
        Create the standard 7-row metadata table that appears at the top
        of every DOCX storyboard template.

        Table layout (RTL — col 0 on RIGHT):
        +============================================+
        | [TEMPLATE_TITLE]  (merged, teal, white)    |
        +--------------------+-----------------------+
        |    رمز العنصر      | [element_code]        |
        +--------------------+-----------------------+
        |   اسم المشروع      | [project_name]        |
        +--------------------+-----------------------+
        | رقم/اسم الوحدة     | [unit]                |
        +--------------------+-----------------------+
        |   اسم العنصر       | [element_name]        |
        +--------------------+-----------------------+
        | المصمم التعليمي     | [designer]            |
        +--------------------+-----------------------+
        |     التاريخ         | [date]                |
        +--------------------+-----------------------+

        Colors: Teal header (#31849B), light blue labels (#DBE5F1)
        Font: Sakkal Majalla, Bold
        Width: 13950 dxa (col 0: 4050, col 1: 9900)
        Borders: thin black outer (sz=4), thick white inner (sz=18)

        Returns:
            The created Table object.
        """
        # Build the unit string: "الوحدة [number]: [name]"
        unit_str = f"الوحدة {self.unit_number}: {self.unit_name}" if self.unit_name else ""

        # Auto-generate element code if not set
        if not self.element_code:
            unit_padded = str(self.unit_number).zfill(2)
            self.element_code = f"{self.project_code}_U{unit_padded}"

        rows_data = [
            # (label, value, label_shading, value_shading)
            (self.TEMPLATE_TITLE, "", COLOR_HEADER_BG, COLOR_HEADER_BG),  # Row 0: merged title
            ("رمز العنصر", self.element_code, COLOR_LABEL_BG, None),
            ("اسم المشروع", self.project_name, COLOR_LABEL_BG, COLOR_WHITE),
            ("رقم/اسم الوحدة", unit_str, COLOR_LABEL_BG, COLOR_WHITE),
            ("اسم العنصر", self.element_name, COLOR_LABEL_BG, COLOR_WHITE),
            ("المصمم التعليمي", self.designer, COLOR_LABEL_BG, None),
            ("التاريخ", self.date_str, COLOR_LABEL_BG, None),
        ]

        table = self.doc.add_table(rows=len(rows_data), cols=2)
        _set_table_bidi(table)
        _set_table_width(table, META_TABLE_WIDTH)
        _set_table_borders(table, outer_sz=4, inner_sz=18,
                           outer_color=COLOR_BLACK, inner_color=COLOR_WHITE)

        for row_idx, (label, value, label_shading, value_shading) in enumerate(rows_data):
            cell0 = table.cell(row_idx, 0)
            cell1 = table.cell(row_idx, 1)

            _set_cell_width(cell0, META_COL0_WIDTH)
            _set_cell_width(cell1, META_COL1_WIDTH)

            if row_idx == 0:
                # Header row: merged, teal background, white bold text, centered, 14pt
                _merge_cells_in_row(table, 0, 0, 1)
                merged_cell = table.cell(0, 0)
                _set_cell_width(merged_cell, META_TABLE_WIDTH)
                _write_cell(
                    merged_cell, label,
                    font_size_pt=FONT_SIZE_HEADER,
                    bold=True, color_hex=COLOR_HEADER_TEXT,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    shading_hex=label_shading,
                    vertical_alignment="center",
                )
            else:
                # Label cell (col 0): bold, 12pt, black, RTL, light blue background
                _write_cell(
                    cell0, label,
                    font_size_pt=FONT_SIZE_BODY,
                    bold=True, color_hex=COLOR_BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    shading_hex=label_shading,
                    vertical_alignment="center",
                )
                # Value cell (col 1): bold, 12pt, right-aligned
                _write_cell(
                    cell1, value,
                    font_size_pt=FONT_SIZE_BODY,
                    bold=True,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    shading_hex=value_shading,
                    vertical_alignment="center",
                )

        # Apply cell-level border overrides (template uses sz=12 on header, sz=8 on data)
        _set_metadata_cell_borders(table, len(rows_data))

        # Set explicit header row height (template: 1400 twips)
        _set_row_height(table, 0, 1400)

        return table

    def add_arabic_paragraph(self, text, font_size_pt=None, bold=False,
                              color_hex=COLOR_BLACK,
                              alignment=WD_ALIGN_PARAGRAPH.RIGHT):
        """
        Add an Arabic RTL paragraph to the document body.

        Use this for any standalone text between tables, like section
        separators or instructions.

        Args:
            text: Arabic text content.
            font_size_pt: Font size in points (None = default).
            bold: Bold text.
            color_hex: Text color hex.
            alignment: Paragraph alignment.

        Returns:
            The created Paragraph object.
        """
        para = self.doc.add_paragraph()
        para.alignment = alignment
        _set_paragraph_bidi(para)
        _add_rtl_run(para, text, font_size_pt=font_size_pt, bold=bold,
                     color_hex=color_hex)
        return para

    def build_content(self):
        """
        Build the template-specific content tables.

        Subclasses MUST override this method to add their unique tables
        after the metadata table. The base class does nothing.
        """
        pass

    def build(self):
        """
        Build the complete document: header + metadata table + content + footer.

        Call this after setting all content, then call save().

        Returns:
            self (for chaining).
        """
        self._build_header()
        self.create_metadata_table()
        # Add a blank paragraph between metadata and content tables
        self.doc.add_paragraph()
        self.build_content()
        self._add_footer()
        return self

    def save(self, output_path):
        """
        Save the document to disk.

        Creates parent directories if they don't exist.

        Args:
            output_path: Full file path like "output/DSAI/U01/DSAI_U01_Pre_Test.docx"
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)


# =============================================================================
# GROUP A BUILDERS — Infographic-style content table
# (Objectives, Summary, Learning Map share the same table structure)
# =============================================================================

class _GroupABuilder(DocxBuilder):
    """
    Base for Group A templates (Objectives, Summary, Learning Map).

    These templates share an identical content table structure:

    ASCII mockup (RTL — col 0 on RIGHT):
    +==============================================+
    | الشاشة / الانفوجرافيك  (merged, #DBE5F1)     |  <- Row 0
    +-----------+----------------------------------+
    | شاشة      |  (image area)                    |  <- Row 1
    | توضيحية   |                                  |
    +-----------+----------------------------------+
    | النص       |  (content text)                  |  <- Row 2
    | العلمي     |                                  |
    +-----------+----------------------------------+
    | مصادر     |  (image sources)                  |  <- Row 3
    | الصور     |                                  |
    +-----------+----------------------------------+
    | الوصف     |  (description)                    |  <- Row 4
    | التفصيلي  |                                  |
    +-----------+----------------------------------+

    Width: 14175 dxa (col 0: 3015, col 1: 11160)
    All borders: visible black (sz=4)
    Col 0: Bold, CENTER, RTL, #FFFFFF background
    Row 0: merged header, #DBE5F1 background
    """

    # Content for each row in the content table
    CONTENT_TABLE_HEADER = "الشاشة / الانفوجرافيك"
    CONTENT_TABLE_ROWS = [
        "شاشة توضيحية للانفوجرافيك",
        "النص العلمي المعروض على الشاشة",
        "مصادر الصور (إن وجدت(",
        "الوصف التفصيلي للشاشة إن لزم",
    ]

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # Storage for content values (parallel to CONTENT_TABLE_ROWS)
        self._content_values = [""] * len(self.CONTENT_TABLE_ROWS)

    def set_screen_description(self, value):
        """Set the screen/infographic visual description (row 1)."""
        self._content_values[0] = value

    def set_content_text(self, value):
        """Set the main content text displayed on screen (row 2)."""
        self._content_values[1] = value

    def set_image_sources(self, value):
        """Set image sources/credits (row 3)."""
        self._content_values[2] = value

    def set_detailed_description(self, value):
        """Set the detailed screen description (row 4)."""
        self._content_values[3] = value

    def build_content(self):
        """Build the Group A content table."""
        num_rows = 1 + len(self.CONTENT_TABLE_ROWS)  # 1 header + data rows
        table = self.doc.add_table(rows=num_rows, cols=2)
        _set_table_bidi(table)
        _set_table_width(table, GROUP_A_TABLE_WIDTH)
        # Group A uses visible black borders on ALL sides (including inside)
        _set_table_borders(table, outer_sz=4, inner_sz=4,
                           outer_color=COLOR_BLACK, inner_color=COLOR_BLACK)

        # Row 0: merged header — 14pt bold
        _merge_cells_in_row(table, 0, 0, 1)
        header_cell = table.cell(0, 0)
        _set_cell_width(header_cell, GROUP_A_TABLE_WIDTH)
        _write_cell(
            header_cell, self.CONTENT_TABLE_HEADER,
            font_size_pt=FONT_SIZE_HEADER,
            bold=True, color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_LABEL_BG,
        )

        # Data rows
        for i, label in enumerate(self.CONTENT_TABLE_ROWS):
            row_idx = i + 1
            cell0 = table.cell(row_idx, 0)
            cell1 = table.cell(row_idx, 1)

            _set_cell_width(cell0, GROUP_A_COL0_WIDTH)
            _set_cell_width(cell1, GROUP_A_COL1_WIDTH)

            # Template: R1-R2 labels are bold (b+bCs), R3-R4 are bCs-only (not visually bold)
            # vAlign: R1=none, R2=none, R3=center, R4=center
            label_bold = (i <= 1)
            label_valign = "center" if i >= 2 else None
            _write_cell(
                cell0, label,
                font_size_pt=FONT_SIZE_BODY,
                bold=label_bold, color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                shading_hex=COLOR_WHITE,
                vertical_alignment=label_valign,
            )
            _write_cell(
                cell1, self._content_values[i],
                font_size_pt=FONT_SIZE_BODY,
                color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT,
                vertical_alignment=None,
                line_spacing=360,  # 1.5x line spacing for content readability
            )

        return table


class ObjectivesBuilder(_GroupABuilder):
    """
    Builder for Learning Objectives (الأهداف التعليمية) storyboard.

    Uses the Group A infographic-style content table.

    Example:
        builder = ObjectivesBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("الأهداف التعليمية")
        builder.set_element_code("DSAI_U01_MLO")
        builder.set_content_text("1. يتعرف على مفهوم الذكاء الاصطناعي\\n2. ...")
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_MLO.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو إنفوجرافيك"


class SummaryBuilder(_GroupABuilder):
    """
    Builder for Summary (الملخص) storyboard.

    Identical structure to ObjectivesBuilder.

    Example:
        builder = SummaryBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("ملخص الوحدة")
        builder.set_element_code("DSAI_U01_Summary")
        builder.set_content_text("ملخص المحتوى التعليمي...")
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Summary.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو إنفوجرافيك"


class InfographicBuilder(_GroupABuilder):
    """
    Builder for Learning Map / Infographic (خارطة التعلم) storyboard.

    Identical structure to ObjectivesBuilder.

    Example:
        builder = InfographicBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("خارطة التعلم")
        builder.set_element_code("DSAI_U01_Learning_Map")
        builder.set_content_text("خطوات التعلم في هذه الوحدة...")
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Learning_Map.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو إنفوجرافيك"


# =============================================================================
# GROUP B BUILDERS — Card-style content table
# (Discussion and Assignment share the same table structure)
# =============================================================================

class _GroupBBuilder(DocxBuilder):
    """
    Base for Group B templates (Discussion, Assignment).

    These share a card-style content table where ALL label cells have
    the light-blue background (including data rows, not just headers).
    Inside borders are white (invisible), creating a "card" look.

    ASCII mockup (RTL — col 0 on RIGHT):
    +==============================================+
    | [section title]  (merged, #DBE5F1, bold)     |  <- Row 0
    +-----------+----------------------------------+
    | شاشة      | (content)                        |  <- Row 1
    | توضيحية   |                                  |
    +-----------+----------------------------------+
    | النص       | (content)                        |  <- Row 2
    | العلمي     |                                  |
    +-----------+----------------------------------+
    | تعليمات   | (content)                        |  <- Row 3
    | وإرشادات  |                                  |
    +-----------+----------------------------------+
    | الأهداف   | (content)                        |  <- Row 4
    | المرتبطة  |                                  |
    +-----------+----------------------------------+

    Width: 13950 dxa (col 0: 3330, col 1: 10620)
    Outer borders: thin black (sz=4)
    Inside borders: thick white (sz=18) — invisible
    ALL col 0 cells: #DBE5F1 background, BOLD, CENTER
    """

    # Subclasses set these
    SECTION_TITLE = ""
    CONTENT_ROW_LABELS = []
    # Content alignment for col 1 data cells (Discussion uses JUSTIFY, Assignment uses RIGHT)
    CONTENT_ALIGNMENT = WD_ALIGN_PARAGRAPH.RIGHT

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._content_values = {}

    def set_content(self, label_key, value):
        """
        Set content for a specific row by label key.

        Args:
            label_key: The Arabic label text (must match CONTENT_ROW_LABELS).
            value: The content value to put in col 1.
        """
        self._content_values[label_key] = value

    def set_screen_description(self, value):
        """Set screen description (row 1)."""
        self._content_values[self.CONTENT_ROW_LABELS[0]] = value

    def set_content_text(self, value):
        """Set content text (row 2)."""
        self._content_values[self.CONTENT_ROW_LABELS[1]] = value

    def set_instructions(self, value):
        """Set instructions/guidelines (row 3)."""
        self._content_values[self.CONTENT_ROW_LABELS[2]] = value

    def set_related_objectives(self, value):
        """Set related learning objectives (row 4)."""
        self._content_values[self.CONTENT_ROW_LABELS[3]] = value

    def build_content(self):
        """Build the Group B content table."""
        num_rows = 1 + len(self.CONTENT_ROW_LABELS)
        table = self.doc.add_table(rows=num_rows, cols=2)
        _set_table_bidi(table)
        _set_table_width(table, GROUP_B_TABLE_WIDTH)
        # White inside borders for the "card" look
        _set_table_borders(table, outer_sz=4, inner_sz=18,
                           outer_color=COLOR_BLACK, inner_color=COLOR_WHITE)

        # Row 0: merged header — 14pt bold
        _merge_cells_in_row(table, 0, 0, 1)
        header_cell = table.cell(0, 0)
        _set_cell_width(header_cell, GROUP_B_TABLE_WIDTH)
        _write_cell(
            header_cell, self.SECTION_TITLE,
            font_size_pt=FONT_SIZE_HEADER,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_LABEL_BG,
        )

        # Data rows
        for i, label in enumerate(self.CONTENT_ROW_LABELS):
            row_idx = i + 1
            cell0 = table.cell(row_idx, 0)
            cell1 = table.cell(row_idx, 1)

            _set_cell_width(cell0, GROUP_B_COL0_WIDTH)
            _set_cell_width(cell1, GROUP_B_COL1_WIDTH)

            _write_cell(
                cell0, label,
                font_size_pt=FONT_SIZE_BODY,
                bold=True, color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                shading_hex=COLOR_LABEL_BG,
            )
            value = self._content_values.get(label, "")
            # Row 1 (screen description) is CENTER, content rows use class alignment
            content_align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else self.CONTENT_ALIGNMENT
            _write_cell(
                cell1, value,
                font_size_pt=FONT_SIZE_BODY,
                color_hex=COLOR_BLACK,
                alignment=content_align,
                line_spacing=360,  # 1.5x line spacing for content readability
            )

        return table


class DiscussionBuilder(_GroupBBuilder):
    """
    Builder for Discussion (النقاش) storyboard.

    Example:
        builder = DiscussionBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("نقاش الوحدة الأولى")
        builder.set_element_code("DSAI_U01_Discussion")
        builder.set_content_text("ناقش مع زملائك...")
        builder.set_instructions("شارك في النقاش...")
        builder.set_related_objectives("1. يتعرف على...")
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Discussion.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو نقاش"
    SECTION_TITLE = "نقاش 1"
    # Discussion content cells use JUSTIFY alignment (template: jc=both)
    CONTENT_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY
    CONTENT_ROW_LABELS = [
        "شاشة توضيحية للنقاش",
        "النص العلمي المعروض على الشاشة",
        "تعليمات وإرشادات النقاش",
        "الأهداف التعليمية المرتبطة",
    ]


class AssignmentBuilder(_GroupBBuilder):
    """
    Builder for Assignment (الواجب) storyboard.

    Identical table structure to Discussion, just different labels.

    Example:
        builder = AssignmentBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("واجب الوحدة الأولى")
        builder.set_element_code("DSAI_U01_Assignment")
        builder.set_content_text("اكتب مقالة عن...")
        builder.set_instructions("يرجى التسليم خلال أسبوع...")
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Assignment.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو واجب"
    SECTION_TITLE = "واجب 1"
    CONTENT_ROW_LABELS = [
        "شاشة توضيحية للواجب",
        "النص العلمي المعروض على الشاشة",
        "تعليمات وإرشادات الواجب",
        "الأهداف التعليمية المرتبطة",
    ]


# =============================================================================
# GROUP C: TEST BUILDER — 3 tables (metadata + info + questions)
# =============================================================================

class TestBuilder(DocxBuilder):
    """
    Builder for Test storyboards (pre-test, post-test, course exam).

    This template has 3 tables:
    1. Metadata table (standard 7-row)
    2. Test info table (3 rows: header, description, instructions)
    3. Questions table (4 columns, N rows)

    ASCII mockup of the questions table (RTL):
    +===========+==============+================+==================+
    | نص السؤال | بدائل السؤال | الإجابة الصحيحة | رابط/وصف الصور   |
    +-----------+--------------+----------------+------------------+
    | Q1 text   | أ) ...       | ج              | (image desc)     |
    |           | ب) ...       |                |                  |
    |           | ج) ...       |                |                  |
    |           | د) ...       |                |                  |
    +-----------+--------------+----------------+------------------+
    | Q2 text   | ...          | ...            | ...              |
    +-----------+--------------+----------------+------------------+

    Column widths: 3240, 4433, 4050, 2955 dxa
    Header: #DBE5F1, Sakkal Majalla 14pt, CENTER, RTL
    All borders: visible black (sz=4)

    Usage:
        builder = TestBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("الاختبار القبلي")
        builder.set_element_code("DSAI_U01_Pre_Test")
        builder.set_test_info(
            description="الاختبار القبلي للوحدة الأولى",
            instructions="محاولة واحدة فقط"
        )
        builder.add_question(
            question_text="ما هو الذكاء الاصطناعي؟",
            choices="أ) برنامج حاسوبي\\nب) فرع من علوم الحاسب\\nج) لغة برمجة\\nد) نظام تشغيل",
            correct_answer="ب",
            image_description=""
        )
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Pre_Test.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو اختبار"

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._test_description = ""
        self._test_instructions = ""
        self._questions = []

    def set_test_info(self, description, instructions):
        """
        Set the test description and instructions.

        Args:
            description: Arabic test description (e.g. "الاختبار البعدي للوحدة الأولى").
            instructions: Arabic instructions (e.g. "محاولة واحدة، 10 دقائق").
        """
        self._test_description = description
        self._test_instructions = instructions

    def add_question(self, question_text, choices, correct_answer,
                     image_description=""):
        """
        Add a question to the test.

        Args:
            question_text: The question in Arabic.
            choices: Answer choices as a string (newline-separated: "أ) ...\\nب) ...").
            correct_answer: The correct answer letter/text.
            image_description: Optional image link/description.
        """
        self._questions.append({
            "text": question_text,
            "choices": choices,
            "answer": correct_answer,
            "image": image_description,
        })

    def build_content(self):
        """Build the test info table and questions table."""
        self._build_test_info_table()
        self.doc.add_paragraph()  # spacer
        self._build_questions_table()

    def _build_test_info_table(self):
        """
        Build the test information table (3 rows x 2 cols).

        +======================================+
        | معلومات الاختبار (merged, #DBE5F1)    |
        +----------+---------------------------+
        | الوصف    | [description]              |
        +----------+---------------------------+
        | الإرشادات | [instructions]             |
        +----------+---------------------------+
        """
        table = self.doc.add_table(rows=3, cols=2)
        _set_table_bidi(table)
        _set_table_width(table, TEST_INFO_TABLE_WIDTH)
        _set_table_borders(table, outer_sz=4, inner_sz=4,
                           outer_color=COLOR_BLACK, inner_color=COLOR_BLACK)

        # Row 0: merged header — 14pt bold
        _merge_cells_in_row(table, 0, 0, 1)
        header_cell = table.cell(0, 0)
        _set_cell_width(header_cell, TEST_INFO_TABLE_WIDTH)
        _write_cell(
            header_cell, "معلومات الاختبار",
            font_size_pt=FONT_SIZE_HEADER, bold=True, color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_LABEL_BG,
        )

        # Row 1: Description
        info_rows = [
            ("الوصف", self._test_description),
            ("الإرشادات", self._test_instructions),
        ]

        for i, (label, value) in enumerate(info_rows):
            row_idx = i + 1
            cell0 = table.cell(row_idx, 0)
            cell1 = table.cell(row_idx, 1)

            _set_cell_width(cell0, TEST_INFO_COL0_WIDTH)
            _set_cell_width(cell1, TEST_INFO_COL1_WIDTH)

            _write_cell(
                cell0, label,
                font_size_pt=FONT_SIZE_BODY, bold=True, color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                shading_hex=COLOR_LABEL_BG,
            )
            # Add paragraph spacing (template: before=240, after=240)
            _set_paragraph_spacing(cell0.paragraphs[0], before=240, after=240)

            _write_cell(
                cell1, value,
                font_size_pt=FONT_SIZE_BODY, bold=True, color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                line_spacing=360,  # 1.5x for multi-line instructions
            )
            # Add paragraph spacing to value cells too
            _set_paragraph_spacing(cell1.paragraphs[0], before=240, after=240)

        return table

    def _build_questions_table(self):
        """
        Build the questions table (4 cols, header + N question rows).

        Header columns: نص السؤال | بدائل السؤال | الإجابة الصحيحة | رابط/وصف الصور
        """
        num_rows = 1 + max(len(self._questions), 1)  # at least 1 data row
        table = self.doc.add_table(rows=num_rows, cols=4)
        _set_table_bidi(table)
        _set_table_width(table, TEST_Q_TABLE_WIDTH)
        _set_table_borders(table, outer_sz=4, inner_sz=4,
                           outer_color=COLOR_BLACK, inner_color=COLOR_BLACK)

        # Set negative table indent to extend into margins (template: -714 dxa)
        _set_table_indent(table, -714)

        # Header row
        headers = ["نص السؤال", "بدائل السؤال", "الإجابة الصحيحة", "رابط/وصف الصور (إن وجد)"]
        for col_idx, header_text in enumerate(headers):
            cell = table.cell(0, col_idx)
            _set_cell_width(cell, TEST_Q_COL_WIDTHS[col_idx])
            _write_cell(
                cell, header_text,
                font_size_pt=FONT_SIZE_HEADER, color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                shading_hex=COLOR_LABEL_BG,
            )

        # Question rows
        for q_idx, question in enumerate(self._questions):
            row_idx = q_idx + 1

            # Add row if needed (we created at least 1 data row)
            if row_idx >= num_rows:
                # We need to add more rows — add a row to the table XML
                row_element = parse_xml(f'<w:tr {nsdecls("w")}>'
                    + ''.join(f'<w:tc><w:p/></w:tc>' for _ in range(4))
                    + '</w:tr>')
                table._tbl.append(row_element)

            values = [
                question["text"],
                question["choices"],
                question["answer"],
                question["image"],
            ]
            for col_idx, value in enumerate(values):
                cell = table.cell(row_idx, col_idx)
                _set_cell_width(cell, TEST_Q_COL_WIDTHS[col_idx])

                # Col 0 (question text): bold for visual hierarchy
                # Col 2 (correct answer): bold for emphasis
                is_bold = (col_idx == 0 or col_idx == 2)

                _write_cell(
                    cell, value,
                    font_size_pt=FONT_SIZE_BODY,
                    bold=is_bold,
                    color_hex=COLOR_BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    vertical_alignment=None,
                    line_spacing=360,  # 1.5x for readability in question cells
                )
                # Add paragraph spacing to question cells (template: before=240, after=240)
                _set_paragraph_spacing(cell.paragraphs[0], before=240, after=240)

        return table


# =============================================================================
# GROUP C: ACTIVITY BUILDER — Metadata + N scene tables
# =============================================================================

class ActivityBuilder(DocxBuilder):
    """
    Builder for Interactive Activity (نشاط تفاعلي) storyboard.

    This template has a metadata table followed by one or more scene tables.
    Each scene table has 10 rows x 2 cols.

    ASCII mockup of a scene table (RTL):
    +==============================================+
    | المشهد الأول  (merged, #DBE5F1, bold)        |  <- Row 0
    +--------------------+-------------------------+
    | وصف المشهد         | عناصر المشهد             |  <- Row 1 (both #DBE5F1)
    +--------------------+-------------------------+
    | (scene description)| (content area)           |  <- Row 2
    +--------------------+-------------------------+
    | وصف الصور          | (image descriptions)     |  <- Row 3
    +--------------------+-------------------------+
    | وصف موشن جرافيك    | (motion graphic desc)    |  <- Row 4
    +--------------------+-------------------------+
    | مؤثرات صوتية       | (sound effects)          |  <- Row 5
    +--------------------+-------------------------+
    | نص يظهر على الشاشة  | (on-screen text)        |  <- Row 6
    +--------------------+-------------------------+
    | خطوات النشاط       | (activity steps)         |  <- Row 7
    +--------------------+-------------------------+
    | الإجابة الصحيحة     | (correct answer)        |  <- Row 8
    +--------------------+-------------------------+
    | الأزرار بعد نفاذ    | (buttons after attempts)|  <- Row 9
    | المحاولات          |                          |
    +--------------------+-------------------------+

    Width: 13950 dxa (col 0: 4050, col 1: 9900)
    Outer borders: thick (sz=12)
    Inside borders: thick white (sz=18) — invisible
    Rows 0-1: shaded headers (#DBE5F1)

    Usage:
        builder = ActivityBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("النشاط التفاعلي 1.1")
        builder.set_element_code("DSAI_U01_Activity1.1")
        builder.add_scene(
            title="المشهد الأول",
            description="في هذا المشهد يظهر للطالب...",
            elements="النص التالي يظهر على الشاشة...",
            image_desc="صورة توضيحية للمفهوم",
            motion_desc="-",
            sound_effects="-",
            on_screen_text="التغذية الراجعة للإجابة الصحيحة: ...",
            steps="على الطالب اختيار الإجابة الصحيحة...",
            correct_answer="الإجابة الصحيحة هي: ب",
            buttons='زر "مراجعة المحتوى"\nزر "أعد المحاولة"',
        )
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Activity1.1.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو نشاط تفاعلي"

    def _setup_page(self):
        """Override: Activity template uses 1.0cm header distance."""
        super()._setup_page()
        section = self.doc.sections[0]
        section.header_distance = Cm(HEADER_DISTANCE_CM_LARGE)

    # Row labels for the scene table (rows 2-9 data)
    SCENE_ROW_LABELS = [
        "وصف المشهد",             # Row 1 (sub-header col 0)
        "عناصر المشهد",           # Row 1 (sub-header col 1)
    ]
    SCENE_DATA_LABELS = [
        # (label for col 0, key name for data dict)
        ("", "description"),                           # Row 2: scene description
        ("وصف الصور", "image_desc"),                    # Row 3
        ("وصف موشن جرافيك (إن لزم)", "motion_desc"),   # Row 4
        ("مؤثرات صوتية خاصة ( أن لزم )", "sound_effects"),  # Row 5
        ("نص يظهر على الشاشة", "on_screen_text"),       # Row 6
        ("خطوات النشاط", "steps"),                      # Row 7
        ("الإجابة الصحيحة", "correct_answer"),           # Row 8
        ("الأزرار التي تظهر بعد نفاذ المحاولات", "buttons"),  # Row 9
    ]

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._scenes = []

    def add_scene(self, title, description="", elements="",
                  image_desc="-", motion_desc="-", sound_effects="-",
                  on_screen_text="", steps="", correct_answer="",
                  buttons='زر "مراجعة المحتوى"\nزر "أعد المحاولة"'):
        """
        Add a scene to the activity.

        Args:
            title: Scene title (e.g. "المشهد الأول").
            description: Scene description text (goes in row 2 col 0).
            elements: Scene elements content (goes in row 2 col 1).
            image_desc: Image descriptions.
            motion_desc: Motion graphic description.
            sound_effects: Sound effects description.
            on_screen_text: Text shown on screen (includes feedback text).
            steps: Activity steps.
            correct_answer: Correct answer text.
            buttons: Buttons shown after attempts exhausted.
        """
        self._scenes.append({
            "title": title,
            "description": description,
            "elements": elements,
            "image_desc": image_desc,
            "motion_desc": motion_desc,
            "sound_effects": sound_effects,
            "on_screen_text": on_screen_text,
            "steps": steps,
            "correct_answer": correct_answer,
            "buttons": buttons,
        })

    def build_content(self):
        """Build one scene table per scene."""
        for scene in self._scenes:
            self._build_scene_table(scene)
            self.doc.add_paragraph()  # spacer between scenes

    def _build_scene_table(self, scene):
        """Build a single 10-row scene table."""
        table = self.doc.add_table(rows=10, cols=2)
        _set_table_bidi(table)
        _set_table_width(table, ACTIVITY_TABLE_WIDTH)
        # Activity uses THICK outer borders (sz=12) and white inner borders
        _set_table_borders(table, outer_sz=12, inner_sz=18,
                           outer_color=COLOR_BLACK, inner_color=COLOR_WHITE)

        # Row 0: merged scene title header — 14pt bold
        _merge_cells_in_row(table, 0, 0, 1)
        header_cell = table.cell(0, 0)
        _set_cell_width(header_cell, ACTIVITY_TABLE_WIDTH)
        _write_cell(
            header_cell, scene["title"],
            font_size_pt=FONT_SIZE_HEADER,
            bold=True, color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_LABEL_BG,
            vertical_alignment="center",
        )

        # Row 1: sub-headers (both shaded) — 12pt bold
        cell0 = table.cell(1, 0)
        cell1 = table.cell(1, 1)
        _set_cell_width(cell0, ACTIVITY_COL0_WIDTH)
        _set_cell_width(cell1, ACTIVITY_COL1_WIDTH)
        _write_cell(
            cell0, "وصف المشهد",
            font_size_pt=FONT_SIZE_BODY,
            bold=True, color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_LABEL_BG,
            vertical_alignment="center",
        )
        _write_cell(
            cell1, "عناصر المشهد",
            font_size_pt=FONT_SIZE_BODY,
            bold=True, color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_LABEL_BG,
            vertical_alignment="center",
        )

        # Rows 2-9: data rows
        data_rows = [
            (scene["description"], scene["elements"]),
            ("وصف الصور", scene["image_desc"]),
            ("وصف موشن جرافيك (إن لزم)", scene["motion_desc"]),
            ("مؤثرات صوتية خاصة ( أن لزم )", scene["sound_effects"]),
            ("نص يظهر على الشاشة", scene["on_screen_text"]),
            ("خطوات النشاط", scene["steps"]),
            ("الإجابة الصحيحة", scene["correct_answer"]),
            ("الأزرار التي تظهر بعد نفاذ المحاولات", scene["buttons"]),
        ]

        for i, (col0_text, col1_text) in enumerate(data_rows):
            row_idx = i + 2
            cell0 = table.cell(row_idx, 0)
            cell1 = table.cell(row_idx, 1)

            _set_cell_width(cell0, ACTIVITY_COL0_WIDTH)
            _set_cell_width(cell1, ACTIVITY_COL1_WIDTH)

            # Row 2 C0 (description) should be bold per template
            col0_bold = (i == 0)
            # Label cells (rows 3+) are bold for visual hierarchy
            if i >= 1:
                col0_bold = True
            # Rows 7-9 label cells have explicit white shading in template
            col0_shading = COLOR_WHITE if i >= 5 else None

            _write_cell(
                cell0, col0_text,
                font_size_pt=FONT_SIZE_BODY,
                bold=col0_bold,
                color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                shading_hex=col0_shading,
                vertical_alignment=None,
            )

            # Row 6 C1 (on_screen_text) and Row 7 C1 (steps) should be bold
            col1_bold = (i == 4 or i == 5)  # i=4 is "نص يظهر على الشاشة", i=5 is "خطوات النشاط"

            _write_cell(
                cell1, col1_text,
                font_size_pt=FONT_SIZE_BODY,
                bold=col1_bold,
                color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                vertical_alignment=None,
                line_spacing=360,  # 1.5x for content readability
            )

        return table


# =============================================================================
# GROUP C: VIDEO BUILDER — 6-row metadata + N 4-col scene tables
# =============================================================================

class VideoBuilder(DocxBuilder):
    """
    Builder for Motion Video (فيديو موشن) storyboard.

    This is the most complex DOCX template:
    - A 6-row metadata table (NOT 7 — missing "رقم/اسم الوحدة")
    - N scene tables, each with 7 rows x 4 columns

    ASCII mockup of a scene table (RTL):
    +===========================================================+
    |  مشهد العنوان  (merged 4 cols, #CFE2F3 light blue)        |  <- Row 0
    +-----------+-----------------------------------------------+
    | شاشة      |  (image area, merged 3 cols)                  |  <- Row 1
    | توضيحية   |                                               |
    +-----------+-----------------------------------------------+
    | مؤثرات    |  (audio description, merged 3 cols)           |  <- Row 2
    | صوتية     |                                               |
    +-----------+-----------+---------------+-------------------+
    | النص      | النصوص    | الوصف          | روابط             |  <- Row 3 (sub-headers)
    | العلمي    | المعروضة  | التفصيلي       | الصور             |
    | المقروء   |           |                |                   |
    +-----------+-----------+---------------+-------------------+
    | (narr.)   | (on-scrn) | (description)  | (image links)    |  <- Row 4-6 (data)
    +-----------+-----------+---------------+-------------------+

    Column widths: 3490, 3002, 4418, 3050 dxa
    Total width: 13960 dxa
    Scene header: #CFE2F3 (light blue, unique to this template)
    ALL borders: visible black (sz=8, thicker than standard)

    Usage:
        builder = VideoBuilder(
            project_code="DSAI", unit_number=1,
            unit_name="المهارات الرقمية",
            project_name="تطوير 15 مقرر إلكتروني",
            institution="جامعة نجران", designer="أحمد"
        )
        builder.set_element_name("فيديو موشن الوحدة 1")
        builder.set_element_code("DSAI_U01_Video")
        builder.add_scene(
            title="مشهد العنوان",
            screen_description="",
            sound_effects="موسيقى هادئة",
            narration_segments=[
                {
                    "narration": "مرحبا بكم في...",
                    "on_screen_text": "عنوان الفيديو",
                    "scene_description": "يظهر العنوان...",
                    "image_links": "",
                },
            ]
        )
        builder.build()
        builder.save("output/DSAI/U01/DSAI_U01_Video.docx")
    """
    TEMPLATE_TITLE = "قالب سيناريو فيديوهات موشن جرافيك"

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._scenes = []

    def _setup_page(self):
        """Override: Video template uses 1.0cm header distance."""
        super()._setup_page()
        section = self.doc.sections[0]
        section.header_distance = Cm(HEADER_DISTANCE_CM_LARGE)

    def create_metadata_table(self):
        """
        Override: Video template has a 6-row metadata table (missing unit row).

        Also: labels are NOT bold (unlike other templates), and header uses
        14pt font explicitly.
        """
        rows_data = [
            (self.TEMPLATE_TITLE, "", COLOR_HEADER_BG, COLOR_HEADER_BG),
            ("رمز العنصر", self.element_code, COLOR_LABEL_BG, COLOR_WHITE),
            ("اسم المشروع", self.project_name, COLOR_LABEL_BG, COLOR_WHITE),
            ("اسم العنصر", self.element_name, COLOR_LABEL_BG, COLOR_WHITE),
            ("المصمم التعليمي", self.designer, COLOR_LABEL_BG, None),
            ("التاريخ", self.date_str, COLOR_LABEL_BG, COLOR_WHITE),
        ]

        table = self.doc.add_table(rows=len(rows_data), cols=2)
        _set_table_bidi(table)
        _set_table_width(table, META_TABLE_WIDTH)
        _set_table_borders(table, outer_sz=4, inner_sz=18,
                           outer_color=COLOR_BLACK, inner_color=COLOR_WHITE)

        for row_idx, (label, value, label_shading, value_shading) in enumerate(rows_data):
            cell0 = table.cell(row_idx, 0)
            cell1 = table.cell(row_idx, 1)
            _set_cell_width(cell0, META_COL0_WIDTH)
            _set_cell_width(cell1, META_COL1_WIDTH)

            if row_idx == 0:
                _merge_cells_in_row(table, 0, 0, 1)
                merged_cell = table.cell(0, 0)
                _set_cell_width(merged_cell, META_TABLE_WIDTH)
                _write_cell(
                    merged_cell, label,
                    font_size_pt=FONT_SIZE_HEADER,
                    bold=True, color_hex=COLOR_HEADER_TEXT,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    shading_hex=label_shading,
                    vertical_alignment="center",
                )
            else:
                # Video template: labels are NOT bold, but use explicit 12pt
                _write_cell(
                    cell0, label,
                    font_size_pt=FONT_SIZE_BODY,
                    color_hex=COLOR_BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    shading_hex=label_shading,
                    vertical_alignment="center",
                )
                _write_cell(
                    cell1, value,
                    font_size_pt=FONT_SIZE_BODY,
                    color_hex=COLOR_BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    shading_hex=value_shading,
                    vertical_alignment="center",
                )

        # Video metadata uses cell-level sz=8 borders on ALL cells (including header)
        border_8 = {"sz": "8", "val": "single", "color": "000000"}
        for row_idx in range(len(rows_data)):
            for col_idx in range(len(table.columns)):
                try:
                    cell = table.cell(row_idx, col_idx)
                    docx_set_cell_borders(
                        cell,
                        top=border_8, bottom=border_8,
                        left=border_8, right=border_8,
                    )
                except Exception:
                    continue

        # Set header row height (template: 1898 twips for video)
        _set_row_height(table, 0, 1898)

        return table

    def add_scene(self, title, screen_description="", sound_effects="",
                  narration_segments=None):
        """
        Add a scene to the video storyboard.

        Each scene has:
        - A title (e.g. "المشهد الأول")
        - Screen description and sound effects (full-width rows)
        - Multiple narration segments (each fills rows 4-6 in the 4-col grid)

        Args:
            title: Scene title in Arabic.
            screen_description: Visual description for the scene.
            sound_effects: Special sound effects.
            narration_segments: List of dicts, each with keys:
                - "narration": The narrated/read text
                - "on_screen_text": Text shown on screen
                - "scene_description": Detailed scene description
                - "image_links": Image source links/descriptions

        Example narration_segments:
            [
                {
                    "narration": "مرحبا بكم...",
                    "on_screen_text": "العنوان الرئيسي",
                    "scene_description": "يظهر العنوان مع انيميشن...",
                    "image_links": "logo.png",
                },
                {
                    "narration": "في هذا الفيديو سنتعلم...",
                    "on_screen_text": "أهداف الفيديو",
                    "scene_description": "تظهر قائمة بالأهداف...",
                    "image_links": "",
                },
            ]
        """
        if narration_segments is None:
            narration_segments = [
                {"narration": "", "on_screen_text": "", "scene_description": "", "image_links": ""}
            ]
        self._scenes.append({
            "title": title,
            "screen_description": screen_description,
            "sound_effects": sound_effects,
            "segments": narration_segments,
        })

    def build_content(self):
        """Build one scene table per scene."""
        for scene in self._scenes:
            self._build_scene_table(scene)
            self.doc.add_paragraph()  # spacer

    def _build_scene_table(self, scene):
        """
        Build a single 7-row x 4-col scene table for the video template.

        Row structure:
        - Row 0: Scene title (merged 4 cols, #CFE2F3)
        - Row 1: Screen description label + content area (merged 3 cols)
        - Row 2: Sound effects label + content area (merged 3 cols)
        - Row 3: 4 sub-headers for the narration grid
        - Rows 4+: narration data rows (one per segment)
        """
        # Calculate rows: 4 fixed rows + number of segments (min 3 data rows)
        num_data_rows = max(len(scene["segments"]), 3)
        num_rows = 4 + num_data_rows

        table = self.doc.add_table(rows=num_rows, cols=4)
        _set_table_bidi(table)
        _set_table_width(table, VIDEO_TABLE_WIDTH)
        # Video uses sz=8 borders (thicker than standard)
        _set_table_borders(table, outer_sz=8, inner_sz=8,
                           outer_color=COLOR_BLACK, inner_color=COLOR_BLACK)

        # Row 0: Scene title (merged across 4 cols) — 14pt bold
        _merge_cells_in_row(table, 0, 0, 3)
        title_cell = table.cell(0, 0)
        _set_cell_width(title_cell, VIDEO_TABLE_WIDTH)
        _write_cell(
            title_cell, scene["title"],
            font_size_pt=FONT_SIZE_HEADER,
            bold=True,
            color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            shading_hex=COLOR_VIDEO_SCENE,
        )

        # Row 1: Screen description (col 0 = label, cols 1-3 merged)
        cell0 = table.cell(1, 0)
        _set_cell_width(cell0, VIDEO_COL_WIDTHS[0])
        _write_cell(
            cell0, "شاشة توضيحية للمشهد",
            font_size_pt=FONT_SIZE_BODY,
            bold=True,
            color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _merge_cells_in_row(table, 1, 1, 3)
        merged_cell = table.cell(1, 1)
        _write_cell(
            merged_cell, scene["screen_description"],
            font_size_pt=FONT_SIZE_BODY,
            color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        # Row 2: Sound effects (col 0 = label, cols 1-3 merged)
        cell0 = table.cell(2, 0)
        _set_cell_width(cell0, VIDEO_COL_WIDTHS[0])
        _write_cell(
            cell0, "مؤثرات صوتية خاصة",
            font_size_pt=FONT_SIZE_BODY,
            bold=True,
            color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _merge_cells_in_row(table, 2, 1, 3)
        merged_cell = table.cell(2, 1)
        _write_cell(
            merged_cell, scene["sound_effects"],
            font_size_pt=FONT_SIZE_BODY,
            color_hex=COLOR_BLACK,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        # Row 3: Sub-headers for the 4-column narration grid — 12pt bold
        sub_headers = [
            "النص العلمي المقروء",
            "النصوص التي تظهر في المشاهد",
            "الوصف التفصيلي للمشهد والتزامن مع النص المقروء والصور",
            "روابط الصور",
        ]
        for col_idx, header_text in enumerate(sub_headers):
            cell = table.cell(3, col_idx)
            _set_cell_width(cell, VIDEO_COL_WIDTHS[col_idx])
            _write_cell(
                cell, header_text,
                font_size_pt=FONT_SIZE_BODY,
                bold=True,
                color_hex=COLOR_BLACK,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                shading_hex=COLOR_LABEL_BG,
            )

        # Data rows (one per narration segment)
        for seg_idx, segment in enumerate(scene["segments"]):
            row_idx = 4 + seg_idx
            if row_idx >= num_rows:
                break  # safety

            values = [
                segment.get("narration", ""),
                segment.get("on_screen_text", ""),
                segment.get("scene_description", ""),
                segment.get("image_links", ""),
            ]
            for col_idx, value in enumerate(values):
                cell = table.cell(row_idx, col_idx)
                _set_cell_width(cell, VIDEO_COL_WIDTHS[col_idx])

                # Col 0 (narration text): italic to distinguish from descriptions
                is_narration = (col_idx == 0)

                _write_cell(
                    cell, value,
                    font_size_pt=FONT_SIZE_BODY,
                    bold=is_narration,  # Narration text bold for emphasis
                    color_hex=COLOR_BLACK,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                    vertical_alignment=None,
                    line_spacing=360,  # 1.5x for readability in data cells
                )

        return table
