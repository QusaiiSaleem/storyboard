"""
RTL Helpers — Shared workarounds for Arabic text in python-pptx and python-docx
================================================================================

Neither python-pptx nor python-docx has complete built-in RTL support.
Both require XML-level workarounds for proper Arabic text rendering.

This module provides all RTL/bidi helper functions used by both the
PPTX engine (LectureBuilder) and the DOCX engine (DocxBuilder).

Why this module exists:
    python-pptx doesn't expose an RTL property on paragraphs — you have
    to set it directly on the XML. Similarly, python-pptx only sets the
    "latin" font when you use font.name, but Arabic text needs the
    "cs" (Complex Script) font to be set separately via XML.

Usage:
    from engine.rtl_helpers import (
        pptx_set_paragraph_rtl,
        pptx_set_run_font_arabic,
        pptx_set_paragraph_ltr,
    )
"""

from lxml import etree

# ============================================================================
# python-pptx helpers — for Interactive Lecture presentations
# ============================================================================

# The XML namespace for DrawingML (used in PPTX files)
_DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def pptx_set_paragraph_rtl(paragraph):
    """
    Set paragraph direction to RTL for Arabic text in PowerPoint.

    In PowerPoint's XML, the paragraph properties element is <a:pPr>
    and the RTL attribute is rtl="1". This function adds or updates
    that attribute.

    Args:
        paragraph: A python-pptx paragraph object

    Example:
        >>> from pptx.util import Pt
        >>> p = text_frame.paragraphs[0]
        >>> run = p.add_run()
        >>> run.text = "مرحبا"
        >>> pptx_set_paragraph_rtl(p)  # Now the text flows right-to-left
    """
    pPr = paragraph._p.get_or_add_pPr()
    pPr.set('rtl', '1')


def pptx_set_paragraph_ltr(paragraph):
    """
    Set paragraph direction to LTR (used for slide numbers, etc.)

    Even in an Arabic RTL presentation, numbers and some labels
    should still flow left-to-right.

    Args:
        paragraph: A python-pptx paragraph object
    """
    pPr = paragraph._p.get_or_add_pPr()
    pPr.set('rtl', '0')


def pptx_set_run_font_arabic(run, font_name, language="ar-JO"):
    """
    Set the font for Arabic text on a python-pptx run.

    python-pptx's font.name only sets the "latin" font. For Arabic
    text, PowerPoint uses the "cs" (Complex Script) font. This function
    sets ALL THREE font slots (cs, latin, ea) to ensure the correct
    font renders regardless of which engine PowerPoint uses.

    It also sets the language tag so PowerPoint applies proper Arabic
    text shaping rules.

    Args:
        run: A python-pptx Run object
        font_name: The font family name (e.g., "Tajawal ExtraBold")
        language: Language tag for text shaping (default: "ar-JO")

    Why we set all three font slots:
        - cs (Complex Script): Used for Arabic, Hebrew, and other RTL scripts
        - latin: Used for Latin/ASCII characters mixed into Arabic text
        - ea (East Asian): Set for completeness — prevents fallback issues

    Example:
        >>> run = paragraph.add_run()
        >>> run.text = "مرحبا"
        >>> run.font.size = Pt(18)
        >>> pptx_set_run_font_arabic(run, "Tajawal ExtraBold")
    """
    rPr = run._r.get_or_add_rPr()

    # Complex Script font — this is what Arabic text actually uses
    cs_font = rPr.find(f'{{{_DRAWINGML_NS}}}cs')
    if cs_font is None:
        cs_font = etree.SubElement(rPr, f'{{{_DRAWINGML_NS}}}cs')
    cs_font.set('typeface', font_name)

    # Latin font — used for any English/ASCII characters in the text
    latin_font = rPr.find(f'{{{_DRAWINGML_NS}}}latin')
    if latin_font is None:
        latin_font = etree.SubElement(rPr, f'{{{_DRAWINGML_NS}}}latin')
    latin_font.set('typeface', font_name)

    # East Asian font — set for completeness
    ea_font = rPr.find(f'{{{_DRAWINGML_NS}}}ea')
    if ea_font is None:
        ea_font = etree.SubElement(rPr, f'{{{_DRAWINGML_NS}}}ea')
    ea_font.set('typeface', font_name)

    # Set language tag for proper Arabic text shaping
    rPr.set('lang', language)


# ============================================================================
# python-docx helpers — for DOCX storyboard templates
# ============================================================================
# These will be used by the DOCX engine (docx_engine.py) once it's built.
# The template-analyst is building that module.

def docx_set_paragraph_rtl(paragraph):
    """
    Set paragraph direction to RTL in a Word document.

    Adds <w:bidi/> to the paragraph properties <w:pPr>.
    Required for every paragraph containing Arabic text.

    Args:
        paragraph: A python-docx paragraph object

    Note:
        This uses insert_element_before() to place the bidi element
        in the correct position within the pPr element, as required
        by the Office Open XML schema ordering rules.
    """
    # Import here to avoid requiring python-docx when only using pptx
    from docx.oxml.parser import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.insert_element_before(
        bidi,
        *(
            "w:adjustRightInd",
            "w:snapToGrid",
            "w:spacing",
            "w:ind",
            "w:contextualSpacing",
            "w:mirrorIndents",
            "w:suppressOverlap",
            "w:jc",
            "w:textDirection",
            "w:textAlignment",
            "w:textboxTightWrap",
            "w:outlineLvl",
            "w:divId",
            "w:cnfStyle",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
    )


def docx_set_run_rtl(run):
    """
    Set RTL on a specific run for proper complex script font selection.

    In Word documents, when RTL is set on a run, Word uses the cs_font
    (Complex Script font) instead of the regular font. This is critical
    for Arabic text to render with the correct font.

    Args:
        run: A python-docx Run object
    """
    from docx.oxml import OxmlElement

    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def docx_set_table_rtl(table):
    """
    Set table direction to RTL.

    This reverses the visual column order so that the first column
    appears on the right side (standard for Arabic tables).

    Also sets autofit=False so Word respects explicit column widths.
    Without this, Word may auto-adjust and ignore our DXA widths.

    Uses the built-in API first, with an XML fallback if it fails.

    Args:
        table: A python-docx Table object
    """
    from docx.oxml import OxmlElement

    try:
        from docx.enum.table import WD_TABLE_DIRECTION
        table.table_direction = WD_TABLE_DIRECTION.RTL
    except Exception:
        # Fallback: manipulate XML directly
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)

    # Disable autofit so Word respects our explicit column widths
    table.autofit = False


def docx_set_cell_shading(cell, color_hex):
    """
    Set background color on a table cell.

    IMPORTANT: Creates a NEW shading element each time.
    XML elements get MOVED (not copied) when appended to another parent,
    so you must never reuse the same element across multiple cells.

    Args:
        cell: A python-docx table cell object
        color_hex: Hex color string without # (e.g., "31849B")

    Example:
        >>> docx_set_cell_shading(table.cell(0, 0), "31849B")  # Teal
        >>> docx_set_cell_shading(table.cell(1, 0), "DBE5F1")  # Light blue
    """
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    # CRITICAL: Create a NEW element each time — never reuse!
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def docx_set_cell_borders(cell, **kwargs):
    """
    Set cell borders using XML manipulation.

    python-docx doesn't expose a border API, so we create the XML directly.

    Args:
        cell: A python-docx table cell object
        **kwargs: Border definitions for each edge. Each is a dict with
                  keys like "sz" (size), "val" (style), "color".

    Example:
        >>> # Visible black border on all sides:
        >>> docx_set_cell_borders(cell,
        ...     top={"sz": 4, "val": "single", "color": "000000"},
        ...     bottom={"sz": 4, "val": "single", "color": "000000"},
        ...     start={"sz": 4, "val": "single", "color": "000000"},
        ...     end={"sz": 4, "val": "single", "color": "000000"},
        ... )
        >>> # Invisible white border (thick white = invisible separator):
        >>> docx_set_cell_borders(cell,
        ...     top={"sz": 18, "val": "single", "color": "FFFFFF"},
        ... )
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right', 'start', 'end'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            element = OxmlElement(f'w:{edge}')
            for attr_name, attr_val in edge_data.items():
                element.set(qn(f'w:{attr_name}'), str(attr_val))
            tcBorders.append(element)

    tcPr.append(tcBorders)
