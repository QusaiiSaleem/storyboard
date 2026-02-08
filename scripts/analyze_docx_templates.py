"""
Deep analysis of all DOCX template files.
Extracts: page dimensions, margins, tables, cell formatting, fonts, colors, styles, headers/footers.
"""
import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATES_DIR = "/Users/qusaiabushanap/dev/storyboard/templates"

DOCX_FILES = [
    "قالب الأهداف التعليمية.docx",
    "قالب الاختبارات.docx",
    "قالب الملخص.docx",
    "قالب النشاط.docx",
    "قالب النقاش.docx",
    "قالب الواجب.docx",
    "قالب خارطة التعلم.docx",
    "قالب فيديو.docx",
]

def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 2)

def emu_to_inches(emu):
    if emu is None:
        return None
    return round(emu / 914400, 2)

def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / 12700, 1)

def get_color_hex(color_obj):
    """Extract hex color from various color objects."""
    if color_obj is None:
        return None
    try:
        if hasattr(color_obj, 'rgb') and color_obj.rgb is not None:
            return str(color_obj.rgb)
        if hasattr(color_obj, 'theme_color') and color_obj.theme_color is not None:
            return f"theme:{color_obj.theme_color}"
    except:
        pass
    return None

def get_shading_color(cell):
    """Extract cell shading/background color from XML."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            color = shd.get(qn('w:color'))
            val = shd.get(qn('w:val'))
            return {'fill': fill, 'color': color, 'val': val}
    return None

def get_paragraph_shading(paragraph):
    """Extract paragraph shading from XML."""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            return fill
    return None

def get_cell_borders(cell):
    """Extract cell border details from XML."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        return None
    borders = {}
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        border_el = tcBorders.find(qn(f'w:{side}'))
        if border_el is not None:
            borders[side] = {
                'val': border_el.get(qn('w:val')),
                'sz': border_el.get(qn('w:sz')),
                'color': border_el.get(qn('w:color')),
                'space': border_el.get(qn('w:space')),
            }
    return borders if borders else None

def get_table_borders(table):
    """Extract table-level border details from XML."""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return None
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        return None
    borders = {}
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        border_el = tblBorders.find(qn(f'w:{side}'))
        if border_el is not None:
            borders[side] = {
                'val': border_el.get(qn('w:val')),
                'sz': border_el.get(qn('w:sz')),
                'color': border_el.get(qn('w:color')),
                'space': border_el.get(qn('w:space')),
            }
    return borders if borders else None

def get_cell_vertical_alignment(cell):
    """Get vertical alignment of a cell."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        vAlign = tcPr.find(qn('w:vAlign'))
        if vAlign is not None:
            return vAlign.get(qn('w:val'))
    return None

def get_cell_width(cell):
    """Get cell width from XML."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None:
            w = tcW.get(qn('w:w'))
            wtype = tcW.get(qn('w:type'))
            return {'width': w, 'type': wtype}
    return None

def get_cell_merge_info(cell):
    """Check if cell is merged."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    merge = {}
    vMerge = tcPr.find(qn('w:vMerge'))
    if vMerge is not None:
        val = vMerge.get(qn('w:val'))
        merge['vMerge'] = val if val else 'continue'
    hMerge = tcPr.find(qn('w:hMerge'))
    if hMerge is not None:
        val = hMerge.get(qn('w:val'))
        merge['hMerge'] = val if val else 'continue'
    gridSpan = tcPr.find(qn('w:gridSpan'))
    if gridSpan is not None:
        merge['gridSpan'] = gridSpan.get(qn('w:val'))
    return merge if merge else None

def get_run_details(run):
    """Extract detailed run formatting."""
    details = {
        'text': run.text,
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size_pt': emu_to_pt(run.font.size) if run.font.size else None,
        'font_color': get_color_hex(run.font.color) if run.font.color else None,
    }
    # Check for RTL in run properties
    rPr = run._r.find(qn('w:rPr'))
    if rPr is not None:
        rtl = rPr.find(qn('w:rtl'))
        details['rtl'] = rtl is not None if rtl is not None else None

        # Check for complex script font
        cs = rPr.find(qn('w:rFonts'))
        if cs is not None:
            details['cs_font'] = cs.get(qn('w:cs'))
            details['ascii_font'] = cs.get(qn('w:ascii'))
            details['hAnsi_font'] = cs.get(qn('w:hAnsi'))

        # Check font size for complex script
        szCs = rPr.find(qn('w:szCs'))
        if szCs is not None:
            details['cs_font_size_pt'] = round(int(szCs.get(qn('w:val'))) / 2, 1)

        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            details['font_size_half_pt'] = sz.get(qn('w:val'))

    return details

def get_paragraph_details(para):
    """Extract detailed paragraph formatting."""
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: 'LEFT',
        WD_ALIGN_PARAGRAPH.CENTER: 'CENTER',
        WD_ALIGN_PARAGRAPH.RIGHT: 'RIGHT',
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'JUSTIFY',
        None: None,
    }

    details = {
        'text': para.text,
        'style_name': para.style.name if para.style else None,
        'alignment': alignment_map.get(para.alignment, str(para.alignment)),
        'shading': get_paragraph_shading(para),
        'runs': [],
    }

    # Check for bidi (RTL paragraph direction)
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        bidi = pPr.find(qn('w:bidi'))
        details['bidi_rtl'] = bidi is not None

        # Paragraph spacing
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            details['spacing'] = {
                'before': spacing.get(qn('w:before')),
                'after': spacing.get(qn('w:after')),
                'line': spacing.get(qn('w:line')),
                'lineRule': spacing.get(qn('w:lineRule')),
            }

        # Indentation
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            details['indent'] = {
                'left': ind.get(qn('w:left')),
                'right': ind.get(qn('w:right')),
                'firstLine': ind.get(qn('w:firstLine')),
                'hanging': ind.get(qn('w:hanging')),
            }

    for run in para.runs:
        details['runs'].append(get_run_details(run))

    return details

def analyze_docx(filepath):
    """Full analysis of a DOCX template."""
    doc = Document(filepath)
    analysis = {
        'filename': os.path.basename(filepath),
        'sections': [],
        'tables': [],
        'body_paragraphs': [],
        'styles_used': set(),
    }

    # Analyze sections (page layout)
    for i, section in enumerate(doc.sections):
        sec_info = {
            'section_index': i,
            'page_width_cm': emu_to_cm(section.page_width),
            'page_height_cm': emu_to_cm(section.page_height),
            'page_width_inches': emu_to_inches(section.page_width),
            'page_height_inches': emu_to_inches(section.page_height),
            'margin_top_cm': emu_to_cm(section.top_margin),
            'margin_bottom_cm': emu_to_cm(section.bottom_margin),
            'margin_left_cm': emu_to_cm(section.left_margin),
            'margin_right_cm': emu_to_cm(section.right_margin),
            'header_distance_cm': emu_to_cm(section.header_distance),
            'footer_distance_cm': emu_to_cm(section.footer_distance),
            'orientation': str(section.orientation),
            'different_first_page_header': section.different_first_page_header_footer,
        }

        # Header content
        try:
            header = section.header
            if header and not header.is_linked_to_previous:
                sec_info['header_paragraphs'] = []
                for para in header.paragraphs:
                    sec_info['header_paragraphs'].append(get_paragraph_details(para))
                # Check for header tables
                sec_info['header_tables'] = []
                for table in header.tables:
                    ht = analyze_table(table)
                    sec_info['header_tables'].append(ht)
        except Exception as e:
            sec_info['header_error'] = str(e)

        # Footer content
        try:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                sec_info['footer_paragraphs'] = []
                for para in footer.paragraphs:
                    sec_info['footer_paragraphs'].append(get_paragraph_details(para))
                sec_info['footer_tables'] = []
                for table in footer.tables:
                    ft = analyze_table(table)
                    sec_info['footer_tables'].append(ft)
        except Exception as e:
            sec_info['footer_error'] = str(e)

        analysis['sections'].append(sec_info)

    # Analyze body paragraphs (outside tables)
    for para in doc.paragraphs:
        pdetails = get_paragraph_details(para)
        analysis['body_paragraphs'].append(pdetails)
        if para.style:
            analysis['styles_used'].add(para.style.name)

    # Analyze tables
    for i, table in enumerate(doc.tables):
        analysis['tables'].append(analyze_table(table, i))

    # Convert set to list for JSON serialization
    analysis['styles_used'] = list(analysis['styles_used'])

    return analysis

def analyze_table(table, table_index=0):
    """Analyze a single table in detail."""
    table_info = {
        'table_index': table_index,
        'row_count': len(table.rows),
        'col_count': len(table.columns),
        'table_borders': get_table_borders(table),
        'rows': [],
    }

    # Table alignment
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        jc = tblPr.find(qn('w:jc'))
        if jc is not None:
            table_info['alignment'] = jc.get(qn('w:val'))

        # Table width
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            table_info['table_width'] = {
                'width': tblW.get(qn('w:w')),
                'type': tblW.get(qn('w:type')),
            }

        # Table layout (fixed vs auto)
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is not None:
            table_info['layout'] = tblLayout.get(qn('w:type'))

        # Bidi
        bidi = tblPr.find(qn('w:bidiVisual'))
        table_info['bidi_visual'] = bidi is not None

    # Grid columns
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn('w:gridCol'))
        table_info['grid_columns_emu'] = [gc.get(qn('w:w')) for gc in gridCols]
        table_info['grid_columns_cm'] = [round(int(gc.get(qn('w:w'))) / 914400 * 2.54, 2) if gc.get(qn('w:w')) else None for gc in gridCols]

    for row_idx, row in enumerate(table.rows):
        row_info = {
            'row_index': row_idx,
            'cells': [],
        }

        # Row height
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is not None:
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is not None:
                row_info['height'] = trHeight.get(qn('w:val'))
                row_info['height_rule'] = trHeight.get(qn('w:hRule'))

        for col_idx, cell in enumerate(row.cells):
            cell_info = {
                'row': row_idx,
                'col': col_idx,
                'text': cell.text.strip(),
                'shading': get_shading_color(cell),
                'borders': get_cell_borders(cell),
                'vertical_alignment': get_cell_vertical_alignment(cell),
                'width': get_cell_width(cell),
                'merge_info': get_cell_merge_info(cell),
                'paragraphs': [],
            }

            for para in cell.paragraphs:
                cell_info['paragraphs'].append(get_paragraph_details(para))

            row_info['cells'].append(cell_info)

        table_info['rows'].append(row_info)

    return table_info

def main():
    all_analyses = {}

    for fname in DOCX_FILES:
        filepath = os.path.join(TEMPLATES_DIR, fname)
        if os.path.exists(filepath):
            print(f"\n{'='*60}")
            print(f"Analyzing: {fname}")
            print(f"{'='*60}")
            analysis = analyze_docx(filepath)
            all_analyses[fname] = analysis

            # Print summary
            print(f"  Sections: {len(analysis['sections'])}")
            print(f"  Tables: {len(analysis['tables'])}")
            print(f"  Body paragraphs: {len(analysis['body_paragraphs'])}")
            print(f"  Styles used: {analysis['styles_used']}")

            for sec in analysis['sections']:
                print(f"  Page: {sec['page_width_cm']}cm x {sec['page_height_cm']}cm")
                print(f"  Margins: T={sec['margin_top_cm']}cm B={sec['margin_bottom_cm']}cm L={sec['margin_left_cm']}cm R={sec['margin_right_cm']}cm")

            for t in analysis['tables']:
                print(f"  Table {t['table_index']}: {t['row_count']}rows x {t['col_count']}cols, bidi={t.get('bidi_visual')}")
        else:
            print(f"NOT FOUND: {filepath}")

    # Save full analysis as JSON
    output_path = "/Users/qusaiabushanap/dev/storyboard/docs/template_analysis_raw.json"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_analyses, f, ensure_ascii=False, indent=2, default=str)

    print(f"\n\nFull analysis saved to: {output_path}")
    return all_analyses

if __name__ == "__main__":
    main()
