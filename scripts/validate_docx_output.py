"""
Validate generated DOCX files match the template specs.
Checks page layout, table structure, colors, fonts, RTL.
"""
import sys
sys.path.insert(0, "/Users/qusaiabushanap/dev/storyboard")

from docx import Document
from docx.oxml.ns import qn

OUTPUT_DIR = "/Users/qusaiabushanap/dev/storyboard/output/test"

def emu_to_cm(emu):
    return round(emu / 914400 * 2.54, 2)

def validate_file(filepath, expected_tables, expected_title):
    print(f"\nValidating: {filepath}")
    doc = Document(filepath)
    errors = []

    # Check page layout
    sec = doc.sections[0]
    w = emu_to_cm(sec.page_width)
    h = emu_to_cm(sec.page_height)
    if abs(w - 29.7) > 0.1:
        errors.append(f"Page width {w}cm != 29.7cm")
    if abs(h - 21.0) > 0.1:
        errors.append(f"Page height {h}cm != 21.0cm")
    if emu_to_cm(sec.top_margin) != 2.54:
        errors.append(f"Top margin {emu_to_cm(sec.top_margin)}cm != 2.54cm")

    # Check table count
    if len(doc.tables) != expected_tables:
        errors.append(f"Table count {len(doc.tables)} != expected {expected_tables}")

    # Check first table (metadata)
    if doc.tables:
        t0 = doc.tables[0]
        # Check BiDi
        tblPr = t0._tbl.find(qn('w:tblPr'))
        bidi = tblPr.find(qn('w:bidiVisual')) if tblPr else None
        if bidi is None:
            errors.append("Metadata table missing bidiVisual (RTL)")

        # Check header row shading
        header_cell = t0.cell(0, 0)
        tcPr = header_cell._tc.find(qn('w:tcPr'))
        if tcPr:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill and fill.upper() != "31849B":
                    errors.append(f"Header fill {fill} != 31849B")
            else:
                errors.append("Header cell missing shading")

        # Check title text
        title_text = t0.cell(0, 0).text.strip()
        if expected_title and expected_title not in title_text:
            errors.append(f"Title '{title_text[:40]}' doesn't contain '{expected_title[:40]}'")

    # Check footer exists
    footer = sec.footer
    if not footer.paragraphs or not any(p.text.strip() for p in footer.paragraphs):
        # Footer may have field codes that show empty text — check for fldChar
        has_fields = False
        for p in footer.paragraphs:
            for run in p.runs:
                if run._r.find(qn('w:fldChar')) is not None:
                    has_fields = True
                    break
            # Also check child elements for fldChar
            for elem in p._p.iterchildren():
                for sub in elem.iterchildren():
                    if sub.tag == qn('w:fldChar'):
                        has_fields = True
        if not has_fields and not any("Page" in p.text for p in footer.paragraphs):
            errors.append("Footer missing or empty")

    if errors:
        for e in errors:
            print(f"  FAIL: {e}")
    else:
        print(f"  PASS: All checks passed")
    return len(errors) == 0

# Run validations
results = []
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_MLO.docx", 2, "قالب سيناريو إنفوجرافيك"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Summary.docx", 2, "قالب سيناريو إنفوجرافيك"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Learning_Map.docx", 2, "قالب سيناريو إنفوجرافيك"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Discussion.docx", 2, "قالب سيناريو نقاش"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Assignment.docx", 2, "قالب سيناريو واجب"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Pre_Test.docx", 3, "قالب سيناريو اختبار"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Activity1.1.docx", 2, "قالب سيناريو نشاط تفاعلي"))
results.append(validate_file(f"{OUTPUT_DIR}/DSAI_U01_Video.docx", 3, "قالب سيناريو فيديوهات موشن جرافيك"))

print(f"\n{'='*60}")
print(f"Results: {sum(results)}/{len(results)} passed")
print(f"{'='*60}")
